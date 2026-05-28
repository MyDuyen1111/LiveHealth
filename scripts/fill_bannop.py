"""Chèn ảnh + caption vào BanNop.docx → BanNop_filled.docx.

Vị trí chèn:
  4.1.4.2 Thực hiện kiểm thử chức năng (Tìm kiếm)
      → ngay sau paragraph "Kịch bản kiểm thử: Dưới đây là hình ảnh ..."
  4.1.5.2 Thực hiện kiểm thử chức năng Đặt hàng
      → ngay sau paragraph "Kịch bản kiểm thử: Dưới đây là hình ảnh ..."

Captions theo mẫu BanMau:
  "Hình 4. N Kịch bản kiểm thử <chức năng> <tình huống>"
"""
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path("/home/hahoang/Documents/LiveHealth")
SRC = ROOT / "public" / "BanNop.docx"
DST = ROOT / "public" / "BanNop_filled.docx"
IMG_DIR = ROOT / "public" / "anh_baocao"

# (img_filename, caption_after_number)
SECTION_414 = [
    ("4.1.4_TC01_tim_kiem_tu_khoa_hop_le.png",
     "Kịch bản kiểm thử tìm kiếm với từ khóa hợp lệ"),
    ("4.1.4_TC02_tim_kiem_khong_co_ket_qua.png",
     "Kịch bản kiểm thử tìm kiếm không có kết quả"),
    ("4.1.4_TC03_tim_kiem_o_trong.png",
     "Kịch bản kiểm thử tìm kiếm với ô trống"),
]

SECTION_415 = [
    ("4.1.5_TC01_dat_hang_thanh_cong.png",
     "Kịch bản kiểm thử đặt hàng thành công"),
    ("4.1.5_TC02_dat_hang_gio_rong.png",
     "Kịch bản kiểm thử đặt hàng khi giỏ rỗng"),
    ("4.1.5_TC03_dat_hang_chua_dang_nhap.png",
     "Kịch bản kiểm thử đặt hàng khi chưa đăng nhập"),
]

# BanMau numbering convention: 4.1.1→4.1-4.3, 4.1.2→4.4-4.6, 4.1.3→4.7-4.9,
# 4.1.4 bắt đầu Hình 4.10. Ta giữ logic đó cho 4.1.4 (3 hình) và 4.1.5 (3 hình).
START_NUMBER_414 = 10
START_NUMBER_415 = 13


def find_paragraph_idx(doc, predicate):
    """Trả về index paragraph đầu tiên thỏa predicate(text)."""
    for i, p in enumerate(doc.paragraphs):
        if predicate(p.text):
            return i
    return -1


def ensure_caption_style(doc):
    """python-docx không tự tạo style 'Caption' nếu chưa có. Dùng style sẵn có nếu được."""
    for style in doc.styles:
        if style.name == "Caption":
            return style
    return None


def insert_paragraph_after(anchor_para, doc):
    """Tạo một paragraph mới ngay SAU anchor_para (cùng cha)."""
    new_p = OxmlElement("w:p")
    anchor_para._p.addnext(new_p)
    return Paragraph(new_p, anchor_para._parent)


def add_image_run(paragraph, img_path: Path, width_inches=6.0):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(img_path), width=Inches(width_inches))


def add_caption(paragraph, caption_text, caption_style):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption_style is not None:
        paragraph.style = caption_style
    run = paragraph.add_run(caption_text)
    if caption_style is None:
        run.italic = True


def insert_section(doc, anchor_idx, items, start_number, caption_style):
    """Sau anchor_idx, chèn lần lượt: (image, caption) × len(items)."""
    anchor = doc.paragraphs[anchor_idx]
    cursor = anchor  # mỗi phần tử mới đều addnext vào cursor, sau đó cursor ← mới
    for offset, (filename, scenario) in enumerate(items):
        img_path = IMG_DIR / filename
        assert img_path.exists(), f"Missing image: {img_path}"

        # Paragraph chứa ảnh
        img_para = insert_paragraph_after(cursor, doc)
        add_image_run(img_para, img_path, width_inches=6.0)
        cursor = img_para

        # Paragraph caption
        cap_para = insert_paragraph_after(cursor, doc)
        caption_text = f"Hình 4.{start_number + offset} {scenario}"
        add_caption(cap_para, caption_text, caption_style)
        cursor = cap_para


def main():
    print(f"Mở: {SRC}")
    doc = Document(str(SRC))
    caption_style = ensure_caption_style(doc)
    print(f"Caption style: {caption_style.name if caption_style else '(none — fallback italic)'}")

    # Tìm anchor cho 4.1.4 và 4.1.5
    # Cả hai có cùng text "Kịch bản kiểm thử: Dưới đây là hình ảnh ..."
    # Cần lấy ĐÚNG cái nằm sau Heading "4.1.4" và Heading "4.1.5".

    # Lập danh sách (idx, text) các paragraph, rồi quét theo Heading 3 anchor.
    heading_414 = None
    heading_415 = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Heading 3":
            if "4.1.4" in p.text:
                heading_414 = i
            elif "4.1.5" in p.text:
                heading_415 = i

    assert heading_414 is not None, "Không tìm thấy Heading 4.1.4"
    assert heading_415 is not None, "Không tìm thấy Heading 4.1.5"
    print(f"Heading 4.1.4 idx={heading_414}: {doc.paragraphs[heading_414].text}")
    print(f"Heading 4.1.5 idx={heading_415}: {doc.paragraphs[heading_415].text}")

    def find_anchor_between(lo, hi):
        for i in range(lo, hi):
            t = doc.paragraphs[i].text
            if t.startswith("Kịch bản kiểm thử") and "Dưới đây" in t:
                return i
        raise RuntimeError(f"Không tìm thấy anchor 'Kịch bản kiểm thử: Dưới đây ...' trong [{lo},{hi})")

    anchor_414 = find_anchor_between(heading_414, heading_415)
    anchor_415 = find_anchor_between(heading_415, len(doc.paragraphs))
    print(f"Anchor 4.1.4 idx={anchor_414}: {doc.paragraphs[anchor_414].text[:80]}")
    print(f"Anchor 4.1.5 idx={anchor_415}: {doc.paragraphs[anchor_415].text[:80]}")

    # IMPORTANT: chèn 4.1.5 TRƯỚC (anchor index lớn hơn), rồi 4.1.4 — để index 4.1.4 không bị dịch
    insert_section(doc, anchor_415, SECTION_415, START_NUMBER_415, caption_style)
    insert_section(doc, anchor_414, SECTION_414, START_NUMBER_414, caption_style)

    doc.save(str(DST))
    print(f"\nĐã lưu: {DST}")


if __name__ == "__main__":
    main()
